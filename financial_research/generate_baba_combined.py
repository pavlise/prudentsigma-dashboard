from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, datetime

output_dir = r"C:\Users\Pavlos Elpidorou\Documents\AI_Project\financial_research"
os.makedirs(output_dir, exist_ok=True)

TICKER = "BABA"
COMPANY = "Alibaba Group Holding Limited"
DATE = datetime.date.today()

doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

NAVY  = "1A376C"; GOLD = "B8860B"; RED = "C0392B"
GREEN = "1E7E34"; LGRAY = "F4F6FB"; WHITE = "FFFFFF"
DGRAY = "555555"; LYEL = "FFFDE7"; AMBER = "D97706"

def rgb(h): return RGBColor(*bytes.fromhex(h))

def set_bg(cell, h):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),h)
    tcPr.append(shd)

def para(text, bold=False, size=10.5, color=None, italic=False, sb=0, sa=5, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    if align: p.alignment = align
    r = p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb = rgb(color)
    return p

def bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text).font.size = Pt(size); return p

def sh(title, tag=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title.upper()); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=rgb(NAVY)
    if tag:
        r2 = p.add_run(f"  {tag}"); r2.font.size=Pt(9); r2.italic=True; r2.font.color.rgb=rgb(DGRAY)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'1'); b.set(qn('w:color'),NAVY)
    pBdr.append(b); pPr.append(pBdr)

def table(headers, rows, widths=None, hbg=NAVY):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hr = t.rows[0]
    for j,h in enumerate(headers):
        c=hr.cells[j]; c.text=h; set_bg(c,hbg)
        c.paragraphs[0].runs[0].bold=True; c.paragraphs[0].runs[0].font.size=Pt(9.5)
        c.paragraphs[0].runs[0].font.color.rgb=rgb(WHITE)
    for i,rd in enumerate(rows):
        r=t.rows[i+1]; bg=LGRAY if i%2==0 else WHITE
        for j,v in enumerate(rd):
            c=r.cells[j]; c.text=str(v); set_bg(c,bg)
            c.paragraphs[0].runs[0].font.size=Pt(9.5)
    if widths:
        for row in t.rows:
            for j,w in enumerate(widths): row.cells[j].width=Inches(w)
    return t

# ══ COVER ═══════════════════════════════════════════════════════════════════
t = doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before=Pt(0); t.paragraph_format.space_after=Pt(2)
r=t.add_run(f"{COMPANY} ({TICKER})"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=rgb(NAVY)

s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER; s.paragraph_format.space_after=Pt(2)
r2=s.add_run("Combined Deep Research  ·  Value Investing Evaluation  ·  Trading Ideas")
r2.font.size=Pt(11.5); r2.italic=True; r2.font.color.rgb=rgb(DGRAY)

dp=doc.add_paragraph(); dp.alignment=WD_ALIGN_PARAGRAPH.CENTER; dp.paragraph_format.space_after=Pt(14)
r3=dp.add_run(f"Generated: {DATE.strftime('%B %d, %Y')}   |   Exchange: NYSE (ADR)   |   Sector: China Tech / E-Commerce / Cloud AI")
r3.font.size=Pt(9); r3.italic=True; r3.font.color.rgb=rgb("999999")

# ══ 0. DATA TIMELINESS ══════════════════════════════════════════════════════
sh("0. Data Timeliness Confirmation","[Value Investing Rule 1]")
table(["Field","Value"],[
    ("Analysis Date","April 4, 2026"),
    ("Latest Available Annual Report","FY2025 (fiscal year ended March 31, 2025; published May 15, 2025)"),
    ("Most Recent Quarterly Data","Q3 FY2026 (Oct–Dec 2025 quarter; reported ~Feb 2026)"),
    ("Is This the Latest Annual Report?","YES — FY2026 annual (ending March 31, 2026) not yet released"),
    ("Key Caveat","CapEx has surged sharply in Q3 FY2026; FY2026 FCF will differ materially from FY2025"),
    ("Data Reliability","Medium-High — annual confirmed; Q3 trend data requires FY2026 full-year confirmation"),
],widths=[2.8,4.0])
para("⚠ FY2026 FCF WARNING: Q3 FY2026 CapEx reached RMB 31.5B in a single quarter (~$4.4B), driven by "
     "Alibaba's $53B AI/cloud investment commitment. Full-year FY2026 FCF will be significantly lower than "
     "FY2025's $22.87B. All DCF projections incorporate this CapEx cycle explicitly.",
     size=10,color=AMBER,sb=4)

# ══ 1. EXECUTIVE SUMMARY ════════════════════════════════════════════════════
sh("1. Executive Summary","[Trading Ideas]")
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
p.add_run("TACTICAL BUY AT THRESHOLD").bold=True
r_=p.runs[0]; r_.font.size=Pt(13); r_.font.color.rgb=rgb(AMBER)
p.add_run(" — 12-month base-case target $155–$175 (+27–43% from $122.05) | Conviction: MEDIUM-HIGH").font.size=Pt(11)

para("Alibaba trades at $122.05/ADS — exactly at the Graham 30% margin-of-safety buy threshold against a "
     "base-case intrinsic value of ~$174/ADS. The market has aggressively de-rated the stock (-37% from "
     "52-week high of $192.67) on three compounding fears: massive AI/cloud CapEx ($53B over 3 years) "
     "depressing near-term FCF, escalating US-China trade tensions and delisting risk, and e-commerce "
     "market share pressure from PDD/ByteDance. Against these, the bull case is powerful: Alibaba Cloud "
     "growing 34% YoY with AI workloads at triple-digit growth, $49B net cash fortress, $19.3B buyback "
     "remaining, and the Qwen AI model family establishing Alibaba as China's premier AI infrastructure "
     "platform. The risk is real and structural — this is NOT a low-risk, high-margin-of-safety setup "
     "like a deeply undervalued US blue chip. It requires China risk tolerance. For investors who have it, "
     "the risk-reward at $122 is compelling.",size=10.5)

# ══ 2. MARKET SNAPSHOT ══════════════════════════════════════════════════════
sh("2. Market Snapshot (April 4, 2026)","[Deep Research]")
table(["Metric","Value"],[
    ("Current Price (ADR)","$122.05"),
    ("Market Capitalization","~$297B"),
    ("ADSs Outstanding","~2.43B  (1 ADS = 8 ordinary shares)"),
    ("52-Week Range","$95.73 – $192.67"),
    ("Distance from 52-Wk High","-36.7%"),
    ("Distance from 52-Wk Low","+27.5%"),
    ("Trailing P/E","23.35x (on GAAP)"),
    ("FCF Yield (FY2025 FCF)","~7.7%  ($22.87B / $297B)"),
    ("Net Cash Position","~$49B"),
    ("Net Cash Per ADS","~$20.2/ADS"),
    ("EV (Market Cap – Net Cash)","~$248B"),
    ("Revenue Growth (FY2025)","5.3% YoY (cloud accelerating to 34%)"),
    ("Analyst Consensus","STRONG BUY — 16 Buy / 2 Hold / 0 Sell"),
    ("Average Analyst Price Target","$185–$187  (range: $135–$225)"),
    ("Highest PT (J.P. Morgan)","$230"),
],widths=[2.8,3.8])

# ══ 3. BUSINESS OVERVIEW ════════════════════════════════════════════════════
sh("3. Business Overview","[Deep Research]")
para("Alibaba Group is China's dominant technology conglomerate, with six core segments: "
     "(1) Taobao & Tmall Group — China's largest e-commerce ecosystem (Alibaba's primary profit engine via "
     "merchant services, advertising, and commissions; 32% domestic market share); "
     "(2) Cloud Intelligence Group — Asia-Pacific's #1 cloud provider, accelerating at 34% YoY; AI workloads "
     "now >20% of external cloud revenue; Qwen AI model family among global open-source leaders; "
     "(3) International Commerce — AliExpress, Lazada, Trendyol (Turkey), Miravia; high growth but unprofitable; "
     "(4) Cainiao Smart Logistics — intelligent logistics backbone leveraging data; "
     "(5) Local Services — Ele.me (food delivery), Amap (maps); competing with Meituan; "
     "(6) Digital Media & Entertainment — Youku streaming, others.",size=10.5)

# ══ 4. MOAT ANALYSIS ════════════════════════════════════════════════════════
sh("4. Economic Moat Assessment","[Value Investing — Step 2]")
table(["Moat Type","Evidence","Strength"],[
    ("Network Effects (E-Commerce)","Taobao/Tmall: 900M+ annual active consumers × 10M+ merchants; largest buyer-seller marketplace in China","STRONG"),
    ("AI/Cloud Infrastructure Moat","Qwen model family; Alibaba Cloud APAC #1; proprietary AI chips (Hanguang); $53B AI investment building durable lead","MEDIUM→STRONG (building)"),
    ("Brand & Trust (Enterprise)","Aliyun (Alibaba Cloud) trusted by China's largest enterprises; switching costs high once integrated","MEDIUM-STRONG"),
    ("Data Flywheel","Broadest consumer purchase data in China; feeds merchant targeting, credit scoring, AI training","STRONG"),
    ("Logistics / Ecosystem Lock-In","Cainiao integrated with merchant operations; Ele.me riders + Freshippo warehouses enabling 1-hr delivery","MEDIUM"),
    ("Threat — E-Commerce","PDD at 23% market share (vs BABA 32%); ByteDance/Douyin livestream commerce growing fast; price war pressure","SIGNIFICANT"),
    ("Threat — Cloud","Tencent Cloud, Huawei Cloud, ByteDance competing aggressively in domestic cloud","MODERATE"),
],widths=[1.6,4.0,0.9])
para("Overall Moat: MEDIUM-STRONG with widening cloud/AI moat partially offsetting narrowing e-commerce moat. "
     "The $53B CapEx commitment is building a new durable moat in AI infrastructure — but it will take 3–5 "
     "years to manifest in FCF.",size=10.5,color=DGRAY)

# ══ 5. FINANCIAL METRICS ════════════════════════════════════════════════════
sh("5. Key Financial Metrics","[Deep Research + Value Investing Steps 1 & 3]")
table(["Metric","FY2022","FY2023","FY2024","FY2025A","YoY","Source"],[
    ("Revenue","$134.6B","$126.5B","$130.4B","$137.3B","+5.3%","FY2025 annual report"),
    ("Revenue (Cloud only)","~$11B","~$14B","~$18B","~$21B+","+18%+","Quarterly breakdowns"),
    ("GAAP Net Income","$9.7B","$10.5B","$10.7B","$17.9B","+67%","FY2025 annual report"),
    ("Net Margin","7.2%","8.3%","8.2%","13.1%","+490bps","Calculated"),
    ("Gross Margin","43%","37%","37%","40.0%","+300bps","FY2025 annual report"),
    ("Operating Margin","11%","12%","10%","14.1%","+410bps","FY2025 annual report"),
    ("EPS (GAAP, per ADS)","~$4.00","~$4.30","~$4.40","~$5.22","+19%","Derived from P/E"),
    ("FCF (OCF – CapEx)","~$23.5B","~$29.2B","$20.9B","$22.87B","+9.4%","MacroTrends/Earnings"),
    ("FCF Margin","~17%","~23%","~16%","~16.7%","","Calculated"),
    ("Net Cash","—","—","~$45B","~$49B","","Balance sheet"),
    ("Total Payment Volume (GMV)","~$1.3T","~$1.3T","~$1.4T","~$1.5T","+5%","Company disclosures"),
],widths=[1.6,0.75,0.75,0.75,0.75,0.75,1.9])

sh("5b. Q3 FY2026 Quarterly Spotlight (Oct–Dec 2025)","[Most Recent Data]")
table(["Metric","Q3 FY2026","YoY","Commentary"],[
    ("Total Revenue","$38.38B","+8%","Accelerating from FY2025's 5.3% full-year"),
    ("Cloud Intelligence Revenue","Not disclosed","+34%","AI workloads triple-digit growth; 9th consecutive qtr"),
    ("AI as % Cloud Revenue",">20%","Triple-digit growth","Monetization inflection underway"),
    ("CapEx (single quarter)","RMB 31.5B (~$4.4B)","Record high","AI/cloud infra; annualized ~$17-18B"),
    ("Free Cash Flow (Q3 alone)","Deeply negative (RMB -21.8B)","Swung from +13.7B","CapEx outpacing OCF temporarily"),
    ("Taobao App DAU Growth","+20% Aug 2025","Strong","Instant commerce driving engagement"),
],widths=[2.0,1.3,1.0,2.9])
para("Critical observation: FCF turned sharply negative in Q3 FY2026 due to the CapEx surge. This is the "
     "primary reason for the 37% stock selloff from the 52-week high. Parallels exist to Amazon (2013–2016) "
     "and Google (2022–2023) heavy investment cycles that suppressed near-term FCF before monetizing.",
     size=10.5,color=AMBER)

# ══ 6. FCF TREND ════════════════════════════════════════════════════════════
sh("6. FCF Trend & Quality Check","[Value Investing — Step 3]")
table(["Year","FCF","YoY","Commentary"],[
    ("FY2022","~$23.5B","—","Pre-regulation recovery baseline"),
    ("FY2023","~$29.2B","+24%","Peak normalized FCF; minimal CapEx cycle"),
    ("FY2024","$20.9B","-28%","Regulatory hangover + macro slowdown; CapEx picks up"),
    ("FY2025","$22.87B","+9.4%","Recovery; AI CapEx surge beginning"),
    ("FY2026E","$5–10B (est.)","~-60%","Heavy AI/cloud investment year; $17B+ CapEx"),
    ("FY2027E","$10–15B (est.)","Recovery","CapEx levels off; cloud revenues ramp"),
    ("FY2028E","$18–25B (est.)","Growth","Cloud monetization; CapEx moderates per mgmt"),
    ("FY2029E+","$28–40B+ (est.)","Strong growth","Full AI/cloud revenue contribution"),
],widths=[1.0,1.2,0.9,4.1])

table(["Quality Check","Calculation","Verdict"],[
    ("FCF / GAAP Net Income","$22.87B / $17.9B = 1.28x","✓ Strong cash conversion"),
    ("CapEx / Revenue (FY2025)","~$5-6B / $137.3B = ~4%","✓ Moderate — but surging to ~13% in FY2026"),
    ("Current Ratio","1.55x","✓ Adequate liquidity"),
    ("Debt-to-Equity","0.23x","✓ Conservative balance sheet"),
    ("Net Cash per ADS","~$20.2/ADS (16.6% of current price)","✓ Meaningful cash support"),
    ("FCF CAGR (FY2022–FY2025)","($22.87B/$23.5B)^(1/3)–1 = –0.9%","⚠ Essentially flat 3-yr; CapEx-era ahead"),
],widths=[2.2,2.5,2.1])

# ══ 7. ASSUMPTIONS TABLE ════════════════════════════════════════════════════
sh("7. DCF Assumptions & Full Disclosure","[Value Investing Rule 2 — CRITICAL]")
para("A. Hard data from FY2025 annual report (high confidence):",bold=True,size=10.5)
table(["Item","Value","Source"],[
    ("FY2025 Revenue","$137.3B","Alibaba FY2025 annual report (May 15, 2025)"),
    ("FY2025 GAAP Net Income","$17.9B","Annual report — EPS $5.22/ADS"),
    ("FY2025 FCF","$22.87B","MacroTrends / annual cash flow statement"),
    ("Net Cash (FY2025)","~$49B","Balance sheet (cash + ST investments – debt)"),
    ("ADSs Outstanding","~2.43B","Company filings"),
    ("Committed CapEx (3-year)","~$53B ($17.7B/yr average)","Bloomberg / Alibaba announcements Feb 2025"),
    ("Q3 FY2026 CapEx","RMB 31.5B (~$4.4B in 1 quarter)","Q3 FY2026 earnings call"),
],widths=[2.2,1.6,2.8])

para("B. Subjective DCF assumptions — Material uncertainty, read carefully:",bold=True,size=10.5,sb=8)
table(["Parameter","Value","Rationale"],[
    ("WACC — Base","11%","US risk-free 5.5% + equity risk premium 4.5% + China regulatory/VIE risk 1%"),
    ("WACC — Bear","13%","Full China risk pricing: delisting risk, VIE structure failure, antitrust"),
    ("WACC — Bull","9%","Geopolitical risk eases; US-China trade deal; cloud re-rated like AWS"),
    ("CapEx Cycle Assumption","$17B/yr for 3 yrs (FY2026-28), then $8B/yr maintenance","Based on $53B commitment and recent Q3 run-rate"),
    ("FCF Y1-3 (FY2026-28)","$10B→$13B→$17B","OCF ~$27-30B minus elevated CapEx ~$17-18B"),
    ("FCF Y4-5 (FY2029-30)","$24B→$28B","CapEx moderates; cloud revenues ramp significantly"),
    ("FCF Y6-10 (FY2031-35)","12% annual growth","Cloud/AI monetization drives FCF acceleration"),
    ("Terminal Growth Rate","3%","China long-term nominal GDP (~5%) discounted for risk"),
    ("WACC Sensitivity","±1% WACC → ±~22% change in IV","High sensitivity — test both directions"),
    ("FCF growth sensitivity","±3% Y6-10 growth → ±~18% IV change","Significant; cloud monetization key driver"),
    ("China Risk Premium","1–3% above US equivalent","Non-quantifiable; VIE structure, Party influence"),
],widths=[2.0,1.7,2.8])

# ══ 8. DCF VALUATION — BASE CASE ════════════════════════════════════════════
sh("8. DCF Valuation — Base Case (WACC 11%, CapEx Cycle + Recovery)")
para("Starting FCF: $22.87B (FY2025 actual). CapEx surge FY2026-28 depresses FCF. Recovery FY2029+.",
     size=10,color=DGRAY)
table(["Year","FCF (Est.)","Discount Factor (11%)","Present Value"],[
    ("1 — FY2026","$10.0B","0.9009","$9.01B"),
    ("2 — FY2027","$13.0B","0.8116","$10.55B"),
    ("3 — FY2028","$17.0B","0.7312","$12.43B"),
    ("4 — FY2029","$24.0B","0.6587","$15.81B"),
    ("5 — FY2030","$28.0B","0.5935","$16.62B"),
    ("6 — FY2031","$31.4B","0.5346","$16.79B"),
    ("7 — FY2032","$35.2B","0.4817","$16.96B"),
    ("8 — FY2033","$39.4B","0.4339","$17.10B"),
    ("9 — FY2034","$44.1B","0.3909","$17.24B"),
    ("10 — FY2035","$49.4B","0.3522","$17.40B"),
    ("YEARS 1–10 TOTAL","","","$149.91B"),
],widths=[1.5,1.2,2.0,1.9])
table(["Terminal Value Component","Calculation","Result"],[
    ("Year 11 FCF","$49.4B × 1.03","$50.88B"),
    ("Terminal Value (Gordon)","$50.88B ÷ (11% – 3%)","$636.0B"),
    ("PV of Terminal Value","$636.0B ÷ 1.11¹⁰","$223.9B"),
    ("ENTERPRISE VALUE (EV)","$149.91B + $223.9B","$373.8B"),
    ("Add: Net Cash","","+ $49.0B"),
    ("EQUITY VALUE","","$422.8B"),
    ("PER ADS (÷ 2.43B)","","$174.00/ADS"),
    ("Gordon Model Cross-Check","$22.87B×1.08÷(11%–3%)÷2.43B = $127.4B EV → ~$145/ADS","~17% gap — acceptable given CapEx cycle adj."),
],widths=[2.5,2.5,1.7])

# ══ 9. THREE-SCENARIO TABLE ══════════════════════════════════════════════════
sh("9. Three-Scenario Valuation Summary")
table(["Scenario","Core Assumption","WACC","FCF Path","Terminal g","IV/ADS","Upside","Probability"],[
    ("BEAR","VIE failure/delisting; cloud never profits; PDD wins e-comm","13%","$8→10→12B, then 5% growth","2.5%","$78","–36%","20%"),
    ("BASE","CapEx cycle pays off; cloud 34% growth; e-comm stable","11%","$10→28B recovery path","3.0%","$174","+43%","55%"),
    ("BULL","Cloud re-rated like AWS; AI moat established; trade deal","9%","$15→45B, then 12% growth","3.0%","$280","+129%","25%"),
    ("PROB. WEIGHTED","20%×$78 + 55%×$174 + 25%×$280","","","","~$181","+48%","100%"),
],widths=[0.75,2.4,0.6,1.5,0.75,0.8,0.6,0.75])

# ══ 10. MARGIN OF SAFETY ════════════════════════════════════════════════════
sh("10. Margin of Safety — Graham Framework","[Value Investing — Step 5]")
table(["Measure","Value"],[
    ("Base Case Intrinsic Value","$174.00/ADS"),
    ("30% Margin of Safety Buy Price","$121.80/ADS"),
    ("Current Price (April 4, 2026)","$122.05/ADS"),
    ("Current Discount to Base IV","29.9% — AT the 30% threshold"),
    ("Bear-Case Intrinsic Value","$78/ADS"),
    ("Current Premium to Bear Case","+56.5% above bear case (limited downside protection vs bear)"),
    ("Probability-Weighted IV","~$181/ADS"),
    ("Discount to Prob.-Weighted IV","32.6%"),
    ("Graham MoS Threshold (25–35%)","JUST MET on base case; NOT met on bear case"),
    ("Verdict","BUY — but margin is thin; HIGH China risk means bear case must be sized"),
],widths=[3.2,3.4])
para("HONEST ASSESSMENT: Unlike PYPL which was 56% below base-case IV, BABA is only AT the 30% MoS "
     "threshold. The critical difference is BABA's China risk — the bear case at $78 implies a 36% "
     "downside from current prices if geopolitical/VIE risk materializes. The investment requires "
     "deliberate China risk tolerance and appropriate position sizing.",size=10.5,color=AMBER,sb=4)

# ══ 11. PEER COMPARISON ══════════════════════════════════════════════════════
sh("11. Peer Comparison","[Trading Ideas — Fundamental Analysis]")
table(["Company","Price","Mkt Cap","Trailing P/E","Rev Growth","FCF Yield","Cloud Position","Key Risk"],[
    ("Alibaba (BABA)","$122","~$297B","23.4x","5.3%","7.7%","APAC #1 (34% growth)","China/VIE/CapEx"),
    ("JD.com (JD)","~$45","~$70B","14.2x","8%","~4%","Logistics-focused","Thin margins"),
    ("PDD Holdings (PDD)","~$115","~$155B","9.95x","~30%","~8%","None","Regulatory scrutiny"),
    ("Tencent (TCEHY)","~$58","~$780B","~22x","~8%","~3.5%","Competitor","Same China risk"),
    ("Amazon (AMZN)","~$190","~$2T","~35x","11%","~3.5%","AWS #1 Global","Regulatory"),
    ("Microsoft (MSFT)","~$380","~$2.8T","~30x","13%","~2.5%","Azure #2 Global","Antitrust"),
],widths=[1.3,0.7,0.9,0.9,0.9,0.9,1.6,1.1])
para("Valuation Takeaway: BABA at 23.4x trailing P/E looks expensive vs PDD (10x) and JD (14x), but "
     "BABA's cloud business growing 34% deserves a premium over pure-play e-commerce peers. On an "
     "EV/FCF basis (~10.8x based on $248B EV / $22.87B FY2025 FCF), BABA is dramatically cheaper than "
     "AWS-equivalent US cloud comps. The CapEx cycle means FY2026E FCF-based valuation is temporarily "
     "inflated — use normalized FCF or longer-term DCF for fair comparison.",size=10.5)

# ══ 12. CATALYSTS ════════════════════════════════════════════════════════════
sh("12. Catalyst Analysis","[Trading Ideas]")
for label, items in [
    ("Near-Term (0–6 Months)",[
        "FY2026 Full-Year Earnings (est. May 2026) — First full-year report under the heavy CapEx cycle; "
         "any sign cloud revenue growth can absorb CapEx costs = re-rating trigger. Cloud growing 34% YoY.",
        "Q4 FY2026 results will clarify whether AI workload monetization (>20% of external cloud revenue) "
         "is accelerating or stabilizing — the single most important metric to watch.",
        "$19.3B Buyback Remaining (through March 2027) — At $122/ADS, management can retire ~158M ADSs "
         "(6.5% of float) with remaining authorization. Highly accretive signal.",
        "Qwen AI Model Releases — Qwen 3 and subsequent releases competing with GPT-4o/Claude; open-source "
         "adoption drives cloud inference revenue and developer ecosystem stickiness.",
        "US-China Trade Negotiations — Any positive signal on trade tensions could trigger rapid re-rating; "
         "Treasury Secretary Bessent's 'everything on the table' comment (April 2025) was the catalyst for "
         "the recent selloff.",
    ]),
    ("Medium-Term (6–24 Months)",[
        "AI Monetization Inflection — AI workloads >20% of external cloud revenue growing at triple-digit "
         "rates; if this reaches 40-50% within 2 years, BABA Cloud re-rates to AWS/Azure multiples.",
        "Taobao Instant Commerce — Leveraging 2M daily riders and 50K lightning warehouses; competing "
         "directly with Meituan; early data shows 20% DAU growth for Taobao app.",
        "International Commerce Profitability — AliExpress, Lazada, Trendyol scaling toward breakeven; "
         "hidden optionality that isn't valued in current market cap.",
        "CapEx Cycle Moderation (FY2028+) — Management committed to 3-year investment period; if cloud "
         "revenue exceeds investment return expectations, FCF recovery will exceed current forecasts.",
        "Hong Kong Dual Listing — Provides refuge from US delisting risk; HK investors increasingly "
         "active in the float, diversifying the institutional ownership mix.",
    ]),
    ("Event-Driven",[
        "US-China Trade Deal / Tariff Reduction — Single biggest potential positive catalyst; BABA could "
         "rally 30-50% on any meaningful diplomatic progress given it's a key China proxy for global funds.",
        "Stripe/Big Tech AI Partnership — Any Western tech partnership validating Alibaba Cloud/Qwen could "
         "trigger international investor re-engagement.",
        "Jack Ma Rehabilitation — Already partially re-emerged in 2023-2024; any formal government signal "
         "of support for Alibaba's leadership = sentiment catalyst.",
        "M&A by Alibaba — $49B net cash provides firepower to acquire profitable businesses (international "
         "e-commerce, AI startups) that diversify the China risk.",
    ]),
]:
    para(label,bold=True,size=10.5,sb=6)
    for item in items: bullet(item)

# ══ 13. TECHNICAL ANALYSIS ════════════════════════════════════════════════════
sh("13. Technical Context & Options Intelligence","[Trading Ideas]")
table(["Indicator","Reading"],[
    ("Current Price","$122.05 (ADR)"),
    ("52-Week High","$192.67 (likely Jan–Feb 2025 AI/DeepSeek excitement peak)"),
    ("52-Week Low","$95.73"),
    ("Current vs. 52-Wk High","-36.7% — significant correction"),
    ("Current vs. 52-Wk Low","+27.5% — bouncing off lows"),
    ("50-Day Moving Average","$160.1 — stock trading 24% BELOW 50-day MA"),
    ("200-Day Moving Average","$153.4 — stock trading 20% BELOW 200-day MA"),
    ("50-Day vs. 200-Day MA","50-day > 200-day = bullish structural trend (golden cross still intact)"),
    ("RSI","~36–45 (approaching oversold; neutral-to-bearish)"),
    ("Key Support Level 1","$121–$122 (current level; critical floor)"),
    ("Key Support Level 2","$95.73 (52-week low; major support)"),
    ("Resistance Level 1","$130–$135 (20-day EMA zone)"),
    ("Resistance Level 2","$153–$160 (200-day / 50-day MA cluster — major overhead resistance)"),
    ("Resistance Level 3","$192.67 (52-week high; recapturing this = full recovery)"),
    ("Options Intelligence","Elevated put/call ratio reflecting uncertainty; implied vol above historical"),
    ("Overall Technical Outlook","Mixed: structurally bullish (golden cross) but momentum bearish; at support"),
    ("Trading Note","RSI near oversold at major support = potential reversal zone; risk/reward improving"),
],widths=[2.5,4.2])

# ══ 14. INSIDER & INSTITUTIONAL ══════════════════════════════════════════════
sh("14. Insider & Institutional Signals","[Deep Research + Trading Ideas]")
table(["Signal","Detail","Interpretation"],[
    ("Institutional Ownership","~74–76%","High; dominated by global institutional funds"),
    ("US Institutional Holders","1,275 institutions; 26% of total shares held by US institutions","Forced selling risk in delisting scenario"),
    ("Buyback Authorization Remaining","$19.3B through March 2027","Management buying at $122 = conviction signal"),
    ("Recent Buyback (last quarter)","$815M in shares repurchased","Consistent; ~6.7M ADSs at ~$122/ADS"),
    ("Dividend (last paid)","$1.98/ADS (July 10, 2025)","Yield ~1.6% at current price"),
    ("Jack Ma's Stake","~4.8% (reduced from ~6%)","Some selling on recovery from 52-wk low"),
    ("Annual Capital Return (2025)","$22.87B FCF; $3-4B dividends + buybacks","Returning ~25-30% of FCF to shareholders"),
    ("VIE Structure Warning","US investors hold contractual rights, not equity ownership","Structural governance risk; non-quantifiable"),
],widths=[2.1,2.6,1.9])

# ══ 15. RISK ASSESSMENT ══════════════════════════════════════════════════════
sh("15. Risk Assessment","[All Three Frameworks — China Premium Critical]")
para("CATEGORY A — Structural/Existential Risks (affect intrinsic value):",bold=True,size=10.5)
for r in [
    "VIE Structure Risk: US investors own contractual interests, NOT equity in Alibaba's operating entities. If the Chinese government invalidates VIE structures, US shareholders could be left with worthless contracts.",
    "US Delisting Risk (HIGH ALERT): Treasury Secretary Bessent's April 2025 'everything on the table' comment re-ignited fears. If PCAOB audit access is revoked, the 2-year countdown to delisting could begin with year 2 ending in 2026. Forced selling by US institutions (~26% of shares) would be severe.",
    "US Chip Export Controls: NVIDIA H100/H200 export restrictions limit Alibaba's ability to scale GPU infrastructure for AI. This could widen the gap with US hyperscalers. Alibaba's own AI chips (Hanguang 800) partially offset, but not fully.",
    "CapEx Risk — Stranded Investment: $53B committed to AI/cloud infrastructure that may never achieve AWS-like returns if US export controls, competition, or economic slowdown prevent adoption. FCF destruction without revenue payoff is the worst-case scenario.",
    "Regulatory Risk (Chinese Government): Post-2021 crackdown memory remains. Any new antitrust actions, forced divestiture orders, or data security investigations would severely impair the business.",
    "E-Commerce Structural Decline: PDD (Pinduoduo/Temu) growing 30%+ revenue vs BABA's 5%; ByteDance Douyin livestream commerce capturing younger demographics; structural share loss may be irreversible.",
]: bullet(r)

para("CATEGORY B — Macro Risks:",bold=True,size=10.5,sb=8)
for r in [
    "US-China trade war escalation: Additional tariffs on Chinese goods reduce e-commerce volumes on AliExpress/international platforms; sentiment impact on ADR valuation.",
    "China economic slowdown: Consumer confidence and spending directly impact Taobao/Tmall GMV and advertising revenue.",
    "RMB depreciation: Alibaba reports in RMB; USD-denominated ADR investors face currency translation risk.",
    "Global cloud spending slowdown: Enterprise IT budget cuts would disproportionately harm Alibaba Cloud's external revenue growth.",
]: bullet(r)

para("CATEGORY C — Sentiment Only (do not change IV):",bold=True,size=10.5,sb=8)
for r in [
    "Technical momentum: Short-term selling pressure from MA breakdown; does not affect fundamental DCF.",
    "Index rebalancing: MSCI/FTSE weight changes create mechanical selling/buying.",
]: bullet(r)

table(["Risk Factor","Assessment"],[
    ("China/VIE/Geopolitical Risk","HIGH — most important risk; requires explicit position sizing"),
    ("CapEx Cycle Risk","MEDIUM-HIGH — temporary FCF destruction; recoverable if cloud succeeds"),
    ("Competitive Risk (E-Commerce)","MEDIUM — PDD/ByteDance pressure real but BABA stabilizing"),
    ("US Chip Controls (AI)","MEDIUM — partially mitigated by own chip development"),
    ("Beta (vs US market)","~1.4 — higher than market; amplified by China news sensitivity"),
    ("Suggested Position Size","1.5–3% max for risk-tolerant value investors (vs 3–5% for US blue chips)"),
    ("ESG","Governance concerns re: VIE, related-party transactions, Jack Ma history"),
],widths=[2.5,4.2])

# ══ 16. OPERATING RULES ══════════════════════════════════════════════════════
sh("16. Investment Operating Rules","[Value Investing — Step 7]")
para("BUY RULES:",bold=True,size=10.5,color=GREEN)
for r in [
    "Primary entry: Stock ≤ $121.80 (30% MoS vs base IV $174) — CURRENTLY AT THRESHOLD ($122.05).",
    "Execution: Build in 3 tranches: 40% at/near $122, 35% if stock reaches $105–$110, 25% at $95–$100.",
    "Each tranche confirmation required: no new VIE legislation, no PCAOB audit revocation, no new antitrust action.",
    "Monitor Q3/Q4 FY2026 cloud revenue: if growth decelerates below 20% YoY, reduce position.",
    "Combine with Hong Kong-listed BABA (9988.HK) for delisting hedge; HK shares not subject to US delisting risk.",
]: bullet(r)

para("HOLD CONDITIONS (all three must hold):",bold=True,size=10.5,sb=8)
for r in [
    "1. Cloud revenue maintaining >20% YoY growth and FCF trajectory recovering post-FY2028.",
    "2. No VIE structure adverse ruling and no confirmed PCAOB audit deficiency for 2 consecutive years.",
    "3. Management deploying $53B CapEx productively — monitor cloud revenue per dollar of CapEx annually.",
]: bullet(r)

para("SELL TRIGGERS (any single trigger):",bold=True,size=10.5,color=RED,sb=8)
for r in [
    "Price exceeds BULL case intrinsic value of $280/ADS.",
    "PCAOB confirms audit deficiency for 2 consecutive years → imminent delisting risk.",
    "Chinese government initiates new antitrust investigation OR forced divestiture of core segment.",
    "Cloud revenue growth decelerates to <15% YoY for 2 consecutive quarters (CapEx not generating returns).",
    "FCF remains below $10B for 3 consecutive full fiscal years (CapEx permanently destroys value).",
    "Buy thesis on AI/cloud proven wrong: AI workloads decline as % of cloud revenue.",
]: bullet(r)

# ══ 17. RECOMMENDATION SUMMARY ═══════════════════════════════════════════════
sh("17. Recommendation Summary","[All Three Frameworks Combined]")
table(["Metric","Value"],[
    ("Rating","TACTICAL BUY AT THRESHOLD"),
    ("Conviction","MEDIUM-HIGH (with explicit China risk tolerance required)"),
    ("Base-Case Intrinsic Value","$174/ADS"),
    ("Bear-Case Intrinsic Value","$78/ADS (-36% downside)"),
    ("Bull-Case Intrinsic Value","$280/ADS (+129% upside)"),
    ("Probability-Weighted IV","~$181/ADS"),
    ("Current Price","$122.05"),
    ("Discount to Base IV","29.9% — AT the 30% Graham threshold"),
    ("12-Month Trading Target","$155–$175 (technical recovery to MA cluster)"),
    ("Upside to Trading Target","+27–43%"),
    ("12-Month Analyst Avg Target","$185–$187 (+52%)"),
    ("J.P. Morgan Target","$230 (+88%)"),
    ("Key Catalyst to Watch","FY2026 annual results (May 2026) — cloud revenue vs CapEx ROI"),
    ("Position Size (suggested)","1.5–3% max (China risk discount vs US large-cap 3–5%)"),
    ("Hedge Recommendation","Consider HK-listed 9988.HK for US delisting risk hedge"),
    ("Timeframe","24–48 months for full intrinsic value realization"),
    ("FCF Yield (FY2025 normalized)","7.7% — attractive but FY2026 will be far lower"),
    ("Key Risk to Monitor","US delisting / VIE structure validity / PCAOB audit access"),
],widths=[2.8,3.8])

# ══ DISCLAIMER ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
dp2=doc.add_paragraph(); dp2.paragraph_format.space_before=Pt(20)
rd=dp2.add_run(
    "DISCLAIMER: This report is generated for research and informational purposes only. It does not "
    "constitute investment advice, a solicitation to buy or sell securities, or a recommendation of any kind. "
    "Investing in Chinese ADRs involves unique risks including VIE structure uncertainty, US delisting risk, "
    "Chinese regulatory risk, currency risk, and geopolitical risk not present in US-domiciled securities. "
    "All financial data sourced from public filings, earnings releases, and news services as of April 4, 2026. "
    "DCF valuations are highly sensitive to assumptions — see Section 7. Past performance does not guarantee "
    "future results. All investments carry the risk of total loss. Consult a qualified financial advisor "
    "and consider your own risk tolerance before making investment decisions.")
rd.italic=True; rd.font.size=Pt(8.5); rd.font.color.rgb=rgb("999999")

filename = f"{TICKER}_combined_research_{DATE}.docx"
filepath = os.path.join(output_dir, filename)
doc.save(filepath)
print(f"Saved: {filepath}")
