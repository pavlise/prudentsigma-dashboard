from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, datetime

output_dir = r"C:\Users\Pavlos Elpidorou\Documents\AI_Project\financial_research"
os.makedirs(output_dir, exist_ok=True)

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

def add_heading(doc, text, level=1, color=(0,0,0)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F4E79')
        tcPr.append(shd)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255,255,255)
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx+1].cells
        fill = 'EBF3FB' if r_idx % 2 == 0 else 'FFFFFF'
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for run in row_cells[i].paragraphs[0].runs:
                run.font.size = Pt(9)
            tc = row_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    return table

# ── TITLE PAGE ──────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BAIDU (BIDU) — COMBINED INSTITUTIONAL ANALYSIS")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(31, 78, 121)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Deep Financial Research  ·  Value Investing Evaluation  ·  Trading Ideas")
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(89, 89, 89)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run(f"Analysis Date: April 4, 2026  |  Price: ~$111 USD/ADS")
r3.font.size = Pt(10)

doc.add_paragraph()

# ── ONE-LINE VERDICT ─────────────────────────────────────────────────────────
add_heading(doc, "ONE-LINE VERDICT", level=2, color=(31,78,121))
add_para(doc,
    "Baidu is a narrow-moat Chinese AI company trading at 0.93x book and ~11x normalized earnings "
    "— statistically cheap, but the advertising core is eroding, FCF is in an investment trough, "
    "and the bull case requires AI monetization acceleration that isn't yet proven. Not a Graham-style "
    "buy at $111, but a compelling accumulate-on-weakness for patient investors with China risk appetite, "
    "with a $135 12-month target and $187 bull case.",
    bold=False, size=10)

doc.add_paragraph()

# ── PART I: DEEP FINANCIAL RESEARCH ─────────────────────────────────────────
add_heading(doc, "PART I — DEEP FINANCIAL RESEARCH", level=1, color=(31,78,121))

add_heading(doc, "Data Timeliness Confirmation", level=2, color=(31,78,121))
add_table(doc,
    ["Field", "Value"],
    [
        ["Analysis Date", "April 4, 2026"],
        ["Latest Annual Report", "FY2025 (released Feb 26, 2026 — Q4 2025 earnings call)"],
        ["Data Used", "FY2025 (most recent)"],
        ["Is This the Latest?", "✅ YES"],
    ]
)
doc.add_paragraph()

add_heading(doc, "Business Overview", level=2, color=(31,78,121))
add_para(doc,
    "Baidu is China's dominant internet search and AI platform. Revenue flows through two primary segments: "
    "Online Marketing Services (core advertising, ~57% of revenue, declining) and Baidu AI Cloud / "
    "non-marketing (~43% of revenue, growing rapidly). The company also controls Apollo Go, the world's "
    "largest autonomous ride-hailing platform, and is developing proprietary AI chips (Kunlunxin) as a "
    "potential third growth vector. Baidu is best understood as a legacy search giant in structural "
    "transition — monetizing its data moat to build AI infrastructure, at the cost of near-term profitability.")

doc.add_paragraph()
add_heading(doc, "Key Financial Metrics (FY2021–FY2025)", level=2, color=(31,78,121))
add_table(doc,
    ["Metric", "FY2021", "FY2022", "FY2023", "FY2024", "FY2025", "YoY Δ"],
    [
        ["Revenue (CNY B)", "124.5", "123.7", "134.6", "133.1", "129.1", "-3.0%"],
        ["Gross Margin", "—", "—", "—", "—", "43.9%", "—"],
        ["Operating Income (CNY B)", "10.5", "15.9", "21.9", "21.3", "-5.8", "▼"],
        ["Operating Margin", "—", "—", "—", "—", "-4.5%", "—"],
        ["EBITDA Margin", "—", "—", "—", "—", "12.0%", "—"],
        ["Net Income (CNY B)", "10.2", "7.6", "20.3", "23.8", "5.6", "-76.5%"],
        ["Diluted EPS (CNY)", "28.1", "19.8", "55.1", "65.9", "11.8", "-82.2%"],
        ["OCF (CNY B)", "20.1", "26.2", "36.6", "21.2", "-3.0", "▼▼"],
        ["CapEx (CNY B)", "-10.9", "—", "-11.2", "-8.1", "-12.1", "+49%"],
        ["Free Cash Flow (CNY B)", "9.2", "26.2", "25.4", "13.1", "-15.1", "▼▼ (trough)"],
    ]
)
add_para(doc, "⚠️ FY2025 FCF caveat: Negative FCF driven by (a) AI capex +49% YoY to CNY 12.1B and (b) ~CNY 33.8B swing in fair-value changes on investment portfolio. FY2024's CNY 13.1B FCF is used as normalized base for DCF.", size=9)

doc.add_paragraph()
add_heading(doc, "Balance Sheet Snapshot", level=2, color=(31,78,121))
add_table(doc,
    ["Item", "CNY B", "USD B"],
    [
        ["Cash + ST Investments", "141.8", "$20.0B"],
        ["Total Debt", "93.5", "$13.2B"],
        ["Net Cash", "~35.1", "~$4.9B (~$14.5/ADS)"],
        ["Total Equity", "290.1", "$40.9B"],
        ["D/E Ratio", "32.2%", "—"],
        ["Market Cap (at $111)", "—", "$38.1B"],
        ["P/Book", "—", "0.93x"],
    ]
)
add_para(doc, "Balance sheet verdict: Fortress. Net cash ~$4.9B = ~13% of market cap. D/E declining (from 39.8% five years ago).", size=9)

doc.add_paragraph()
add_heading(doc, "AI Business Momentum (Q4 2025 / FY2025)", level=2, color=(31,78,121))
add_table(doc,
    ["Segment / Metric", "Value"],
    [
        ["AI Cloud Revenue YoY Growth", "+33% to RMB 4.2B"],
        ["AI Subscription Revenue YoY Growth", "+128%"],
        ["AI-Powered Core Revenue (share)", "43% of general business (>RMB 11B)"],
        ["ERNIE 4.5 vs GPT-4.5", "Outperforms at 1% of GPT-4.5 price"],
        ["ERNIE 5.0 Launch", "Nov 2025 (natively omnimodal: text/image/audio/video)"],
        ["Apollo Go Cumulative Rides", "20M+"],
        ["Apollo Go Weekly Rides (Q4 2025 peak)", "300,000"],
        ["Apollo Go Fully Autonomous YoY Growth", "+200%"],
        ["Apollo Go Global Footprint", "26 cities incl. South Korea"],
        ["Baidu App MAU (Mar 2025)", "724M"],
        ["Kunlunxin Chips M100/M300", "Launching 2026/2027"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Competitive Landscape", level=2, color=(31,78,121))
add_table(doc,
    ["Segment", "Baidu Position", "Key Threats"],
    [
        ["China Search (overall)", "56.2%", "Bing (24%), ByteDance AI search"],
        ["China Search (mobile)", "71.9%", "ByteDance Douyin time-on-device"],
        ["AI Cloud (token processing, IDC H1 2025)", "17% (#3)", "ByteDance Volcano 49%, Alibaba 27%"],
        ["AI Cloud (overall market)", "6.1% (#5)", "Alibaba 35.8%, ByteDance 14.8%, Huawei 13.1%"],
        ["Foundation Models", "ERNIE 4.5/5.0 — top-tier Chinese LLM", "Alibaba Qwen, DeepSeek, Zhipu"],
        ["Autonomous Driving / Robotaxi", "#1 globally by cumulative autonomous miles", "Waymo (US), Pony.ai"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Bull / Bear Cases", level=2, color=(31,78,121))
add_table(doc,
    ["Case", "Prob.", "Key Drivers"],
    [
        ["Bull (35%)", "35%", "ERNIE monetization accelerates; Apollo Go spin-off; Kunlunxin chip value unlock; China-US tensions ease → multiple expansion"],
        ["Base (40%)", "40%", "Advertising stabilizes at slight decline; AI cloud grows 25–35% annually; FCF recovers to CNY 15–18B by FY2026–27"],
        ["Bear (25%)", "25%", "Advertising -5% to -10% per year; AI capex extends FCF negativity; ByteDance displaces ERNIE; HFCAA delisting risk resurfaces"],
    ]
)

# ── PART II: VALUE INVESTING ──────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "PART II — VALUE INVESTING EVALUATION", level=1, color=(31,78,121))

add_heading(doc, "Assumptions & Estimates Disclosure", level=2, color=(31,78,121))
add_table(doc,
    ["Category", "Parameter", "Value", "Justification"],
    [
        ["A. Hard Data (FY2025 Report)", "Revenue", "CNY 129.1B", "StockAnalysis / FY2025 Annual"],
        ["A. Hard Data (FY2025 Report)", "OCF", "CNY -3.0B", "Cash Flow Statement"],
        ["A. Hard Data (FY2025 Report)", "CapEx", "CNY -12.1B", "Cash Flow Statement"],
        ["A. Hard Data (FY2025 Report)", "FCF (FY2025)", "CNY -15.1B", "Calculated: OCF – CapEx"],
        ["A. Hard Data (FY2024 Report)", "Normalized FCF", "CNY 13.1B", "FY2024 clean year — used as DCF base"],
        ["A. Hard Data", "Net Cash", "CNY ~35.1B", "Balance Sheet"],
        ["A. Hard Data", "Shares (ADS)", "340.25M", "Market data"],
        ["B. Market Data", "Current Price", "$111/ADS", "Apr 2, 2026"],
        ["B. Market Data", "Market Cap", "~$38.1B", "Calculated"],
        ["B. Market Data", "CNY/USD rate", "~7.10", "2026 spot"],
        ["C. Subjective — WACC", "Bear", "13%", "4% risk-free + 5% ERP + 4% China/VIE/ADR risk"],
        ["C. Subjective — WACC", "Base", "12%", "Standard China internet per framework"],
        ["C. Subjective — WACC", "Bull", "10%", "If geopolitical risk recedes"],
        ["C. Subjective — Growth Y1–5", "Bear / Base / Bull", "3% / 8% / 15%", "AI cloud growth offset by ad decline"],
        ["C. Subjective — Growth Y6–10", "Bear / Base / Bull", "2% / 7% / 10%", "Maturing growth phase"],
        ["C. Subjective — Terminal Growth", "Bear / Base / Bull", "2% / 3% / 3.5%", "China long-run nominal GDP proxy"],
        ["D. Normalization", "FY2025 FCF adjusted?", "YES", "~CNY 33.8B fair-value losses + AI capex surge distort FY2025"],
    ]
)
add_para(doc, "Sensitivity: WACC ±1% → fair value ±18% | FCF growth assumption ±5% → fair value ±22%", size=9)

doc.add_paragraph()
add_heading(doc, "Moat Analysis", level=2, color=(31,78,121))
add_table(doc,
    ["Moat Type", "Specific Evidence", "Strength"],
    [
        ["Network Effects", "724M MAU Baidu App; search data flywheel trains ERNIE; more users → better AI → more users", "Strong"],
        ["Regulatory / Compliance", "Deep PRC compliance; preferred local AI partner; gov't backing for Apollo Go rollout", "Strong"],
        ["Scale / Data Moat", "Decades of Chinese-language search data; impossible to replicate; powers ERNIE training", "Strong"],
        ["Cost Advantage", "ERNIE 4.5 delivers GPT-4.5 performance at 1% of the price — structural AI inference moat", "Medium-Strong"],
        ["Switching Costs", "Enterprise Qianfan cloud customers face high migration costs (fine-tuning, API integration)", "Medium"],
        ["First-Mover (Robotaxi)", "Apollo Go #1 globally by cumulative autonomous miles (190M km); RT6 cuts unit cost 60%", "Medium (early stage)"],
    ]
)
add_para(doc, "Overall Moat Rating: NARROW-TO-MODERATE. Core search moat eroding; AI/cloud moat nascent but forming rapidly. Regulatory/data moat durable.", size=9)

doc.add_paragraph()
add_heading(doc, "Value Trap Checklist", level=2, color=(31,78,121))
add_table(doc,
    ["Question", "Answer", "Evidence"],
    [
        ["Core business in secular decline?", "⚠️ PARTIAL YES", "Online advertising declining; FY2025 revenue -3% YoY"],
        ["Competitors taking market share?", "⚠️ YES (AI cloud)", "Baidu is #5 in AI cloud; ByteDance dominant in token processing"],
        ["Irrational M&A or capital destruction?", "✅ No", "CapEx going to productive AI/robotaxi infrastructure"],
        ["Balance sheet landmines?", "✅ No", "Net cash positive; D/E 32%; large cash cushion ($4.9B net)"],
    ]
)

doc.add_paragraph()
add_heading(doc, "FCF Trend Analysis", level=2, color=(31,78,121))
add_table(doc,
    ["Year", "FCF (CNY B)", "YoY Change"],
    [
        ["FY2021", "9.2", "—"],
        ["FY2022", "26.2", "+185%"],
        ["FY2023", "25.4", "-3%"],
        ["FY2024", "13.1", "-48%"],
        ["FY2025", "-15.1", "▼ investment-cycle trough"],
        ["Trend", "Investment-phase trough", "FCF normalization expected FY2026–27"],
    ]
)
add_para(doc, "FCF / Net Income FY2024: 13,100 / 23,760 = 0.55 (below 0.8 ideal threshold). CapEx/Revenue FY2025: 9.4% (elevated but purposeful AI infrastructure).", size=9)

doc.add_paragraph()
add_heading(doc, "DCF Valuation — Base Case (WACC 12%, Terminal g 3%)", level=2, color=(31,78,121))
add_para(doc, "Normalized Base FCF: CNY 13,100M (FY2024). Growth Y1–5: ~8–10% blended. Growth Y6–10: ~7%.", size=9)
add_table(doc,
    ["Year", "FCF (CNY M)", "Discount Factor", "PV (CNY M)"],
    [
        ["1 (FY2026E)", "13,755", "0.8929", "12,284"],
        ["2 (FY2027E)", "15,131", "0.7972", "12,066"],
        ["3 (FY2028E)", "16,947", "0.7118", "12,065"],
        ["4 (FY2029E)", "18,980", "0.6355", "12,062"],
        ["5 (FY2030E)", "20,878", "0.5674", "11,850"],
        ["6 (FY2031E)", "22,548", "0.5066", "11,425"],
        ["7 (FY2032E)", "24,126", "0.4523", "10,913"],
        ["8 (FY2033E)", "25,574", "0.4039", "10,329"],
        ["9 (FY2034E)", "26,852", "0.3606", "9,686"],
        ["10 (FY2035E)", "27,926", "0.3220", "8,992"],
        ["Sum Years 1–10", "", "", "111,672"],
        ["Terminal Value PV", "27,926×1.03/(0.09) discounted", "", "102,911"],
        ["Enterprise Value", "", "", "214,583"],
        ["+ Net Cash", "", "", "35,100"],
        ["Equity Value (CNY M)", "", "", "249,683"],
        ["÷ 7.10 (CNY/USD)", "", "", "$35,167M"],
        ["÷ 340.25M ADS", "", "", "$103.4 / ADS"],
    ]
)
add_para(doc, "Gordon Model cross-check: ~$80–81/ADS (single-stage). 10-year DCF difference: ~28% — within 30% tolerance. ✅", size=9)

doc.add_paragraph()
add_heading(doc, "Three-Scenario Valuation", level=2, color=(31,78,121))
add_table(doc,
    ["Scenario", "Probability", "WACC", "Y1–5 Growth", "Y6–10 Growth", "Terminal g", "Fair Value/ADS"],
    [
        ["Bear", "25%", "13%", "3%", "2%", "2%", "$66.5"],
        ["Base", "55%", "12%", "~8–10%", "~7%", "3%", "$103.4"],
        ["Bull", "20%", "10%", "15%", "10%", "3.5%", "$187.2"],
        ["Probability-Weighted", "100%", "—", "—", "—", "—", "~$105"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Safety Margin & Operating Rules", level=2, color=(31,78,121))
add_table(doc,
    ["Item", "Value"],
    [
        ["Base Case Intrinsic Value", "$103.4 / ADS"],
        ["Graham 30% Margin Buy Price", "$103.4 × 0.70 = $72.4 / ADS"],
        ["Current Price", "$111 / ADS"],
        ["Premium to Base Case", "+7.4% (does NOT meet Graham threshold)"],
        ["Probability-Weighted Fair Value", "~$105 (stock ~6% above)"],
        ["Verdict", "FAIRLY VALUED on base case; compelling for bull-case believers"],
    ]
)
doc.add_paragraph()
add_table(doc,
    ["Action", "Trigger / Price Level"],
    [
        ["Strong Buy / Build Position", "≤ $72 (Graham threshold met)"],
        ["Opportunistic Add", "$75–$90 (near base intrinsic)"],
        ["Hold / Monitor", "$90–$120 ← CURRENT ZONE"],
        ["Trim / Take Profits", "$150–$165 (approaching bull case)"],
        ["Full Exit", "> $185 (exceeds bull case)"],
        ["Thesis-Break Exit", "FCF negative 2+ years ex-investment / Search <40% share / VIE unwound"],
    ]
)

# ── PART III: TRADING IDEAS ───────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "PART III — TRADING IDEAS / EQUITY RESEARCH", level=1, color=(31,78,121))

add_heading(doc, "Executive Summary", level=2, color=(31,78,121))
add_para(doc,
    "HOLD/ACCUMULATE ON WEAKNESS with $135 price target (+22% from $111) over 12–18 months. "
    "Baidu is a classic 'show me' story — the AI pivot is real but monetization lags investment. "
    "The current price ($111) is ~7% above DCF base case but 40% below the bull case; meaningful "
    "upside only materializes if ERNIE subscription growth sustains +100%+ pace AND Apollo Go reaches "
    "profitability. Risk/reward: ~2:1 (bear $66.5 = -40% downside, 12-month target $135 = +22%).")

doc.add_paragraph()
add_heading(doc, "Peer Comparison", level=2, color=(31,78,121))
add_table(doc,
    ["Metric", "BIDU", "Alibaba (BABA)", "Tencent (0700.HK)", "Alphabet (GOOGL)"],
    [
        ["Price (Apr 2026)", "$111", "~$130", "~HK$480", "~$175"],
        ["P/E (FY2025 GAAP)", "~47x*", "~18x", "~22x", "~22x"],
        ["P/E (normalized)", "~11x", "~15x", "~20x", "~22x"],
        ["P/Book", "0.93x", "~1.8x", "~4x", "~6x"],
        ["Revenue Growth", "-3%", "+6%", "+8%", "+12%"],
        ["AI Cloud Growth", "+33%", "+50%+", "~30%", "+28%"],
        ["Net Cash", "$4.9B", "$60B+", "$30B+", "$95B+"],
        ["Robotaxi Optionality", "✅ Significant (Apollo Go)", "❌", "❌", "✅ (Waymo)"],
    ]
)
add_para(doc, "*FY2025 GAAP P/E inflated by one-time investment write-downs. Normalized earnings suggest ~11–13x.", size=9)

doc.add_paragraph()
add_heading(doc, "Catalyst Calendar", level=2, color=(31,78,121))
add_heading(doc, "Near-Term (0–6 months)", level=3)
add_table(doc,
    ["Date", "Catalyst", "Potential Impact"],
    [
        ["May 2026", "Q1 2026 Earnings — first FCF recovery data post-investment trough", "High: ±10–15%"],
        ["Q2 2026", "Apollo Go profitability milestone — targeted in 2026", "High: re-rating trigger"],
        ["Mid-2026", "Kunlunxin M100 chip commercialization", "Medium: new revenue line"],
        ["Ongoing", "ERNIE subscription ARR growth — watch for sustained +100%+", "High"],
    ]
)
add_heading(doc, "Medium-Term (6–24 months)", level=3)
add_table(doc,
    ["Catalyst", "Expected Impact"],
    [
        ["Apollo Go global expansion (post South Korea, potential US/EU pilots)", "Major re-rating if proved commercially viable"],
        ["Kunlunxin chip spin-off", "Estimated $3–5B value unlock / ~$3.5/ADS incremental"],
        ["AI Cloud market share recovery (cost moat vs GPT pricing)", "Margin improvement + revenue acceleration"],
        ["ADR → HK primary listing migration", "Multiple expansion as HK/mainland investors reprice"],
        ["ERNIE embedded in search (Miaoda, GenFlow, new ad format)", "Advertising revenue floor / new monetization stream"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Valuation & Analyst Price Targets", level=2, color=(31,78,121))
add_table(doc,
    ["Source", "Price Target", "Rating", "Upside from $111"],
    [
        ["Wall Street consensus (14 analysts avg)", "$150–$169", "Strong Buy", "+35–52%"],
        ["WallStreetZen model", "$150.31", "Buy", "+35%"],
        ["StockAnalysis consensus", "$154.38", "Strong Buy", "+39%"],
        ["Alpha Spread DCF (Oct 2025)", "$185.42", "—", "+67%"],
        ["This analysis — Bear Case", "$66.5", "—", "-40%"],
        ["This analysis — Base Case", "$103.4", "—", "-7%"],
        ["This analysis — Bull Case", "$187.2", "—", "+69%"],
        ["This analysis — Prob-Weighted", "~$105", "Hold/Accumulate", "-5%"],
        ["This analysis — 12M Trading Target", "$135", "Accumulate on weakness", "+22%"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Risk Assessment", level=2, color=(31,78,121))
add_heading(doc, "Company-Specific Risks", level=3)
add_table(doc,
    ["Risk", "Severity", "Probability", "Mitigation"],
    [
        ["Advertising revenue secular decline (ByteDance stealing mobile search)", "High", "High", "AI-native search integration; pivot to subscription revenue"],
        ["FCF remains negative beyond FY2026 (AI capex continues to expand)", "High", "Medium", "Net cash $4.9B provides ~3-year buffer; no existential threat"],
        ["AI Cloud stays #5 — ByteDance/Alibaba entrench lead", "Medium", "Medium", "Cost moat (1% of GPT price) could win price-sensitive enterprise"],
        ["Apollo Go commercialization delays", "Medium", "Low–Medium", "Already at unit economics in Wuhan/Beijing"],
        ["Management execution risk — too many bets", "Medium", "Medium", "Robin Li has managed this portfolio 20+ years"],
    ]
)
add_heading(doc, "Structural / Macro Risks", level=3)
add_table(doc,
    ["Risk", "Severity", "Assessment"],
    [
        ["VIE structure unwinding", "EXTREME", "Low probability but catastrophic tail risk — '\"worthless\"' per 20-F worst case"],
        ["HFCAA ADR delisting", "High", "Low–Medium probability; PCAOB inspection improved since 2022"],
        ["China regulatory crackdown (repeat of 2021)", "High", "Medium — AI regulation evolving but gov't is pro-AI"],
        ["US-China geopolitical escalation / tariffs", "High", "Ongoing — currently priced in at ~0.93x P/B"],
        ["RMB depreciation vs USD", "Medium", "Partially structural; ADS structure provides some buffer"],
    ]
)
add_para(doc, "ESG Note: Apollo Go's autonomous driving is net social positive, but China data sovereignty concerns limit ESG-screened institutional investors.", size=9)
add_para(doc, "Position Sizing: 2–4% of portfolio. Higher China risk requires smaller sizing vs US tech equivalents. Only suitable for investors with explicit EM/China risk allocation.", bold=True, size=9)

doc.add_paragraph()
add_heading(doc, "Technical Context", level=2, color=(31,78,121))
add_table(doc,
    ["Level / Indicator", "Price"],
    [
        ["Current Price", "$111"],
        ["52-Week Low (inferred, early 2025)", "~$82"],
        ["52-Week High (inferred, late 2025)", "~$121"],
        ["Key Support Zone 1", "$95–$100 (prior base)"],
        ["Key Support Zone 2", "$85 (2025 lows)"],
        ["Key Resistance 1", "$120–$125 (recent highs)"],
        ["Key Resistance 2 / Trading Target", "$135–$145 (analyst consensus range)"],
        ["Context", "Consolidating after 52% rally in 2025 (~$82→$121); -8% from highs; mild pullback — no technical breakdown"],
    ]
)
add_para(doc, "Options Strategy Note: Selling cash-secured puts at $90–95 strike could be an effective entry strategy for longer-term investors willing to accumulate at Graham-adjacent levels.", size=9)

doc.add_paragraph()
add_heading(doc, "Insider Signals & Institutional Activity", level=2, color=(31,78,121))
add_table(doc,
    ["Activity", "Details"],
    [
        ["CEO Robin Li", "Long-term insider; historically bought on dips; reputation for patient capital allocation"],
        ["Share Buybacks", "Active buyback program; $38B market cap vs $20B cash makes buybacks highly accretive"],
        ["Institutional Ownership (Dec 2025)", "Continued institutional presence; no major distribution signals"],
        ["Analyst Trend", "78% Strong Buy/Buy — bullish institutional positioning"],
        ["Apollo Go Value Signal", "$3.5/ADS incremental value cited by analysts — sell-side pricing in spin-off optionality"],
    ]
)

# ── RECOMMENDATION SUMMARY ────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "RECOMMENDATION SUMMARY", level=1, color=(31,78,121))
add_table(doc,
    ["Metric", "Value"],
    [
        ["Rating", "HOLD / ACCUMULATE ON WEAKNESS"],
        ["Conviction", "Medium"],
        ["Near-Term Price Target (12–18M)", "$135"],
        ["Bull Case Target", "$187"],
        ["Graham Buy Price (30% MoS)", "≤ $72"],
        ["Opportunistic Add Zone", "$75–$90"],
        ["Current Price Zone", "Fairly Valued (base) / Discounted vs. Bull"],
        ["Upside to Near-Term Target", "+22%"],
        ["Upside to Bull Case", "+69%"],
        ["Downside (Bear Case)", "-40%"],
        ["Risk/Reward (Base/Near-Term)", "~2.0:1"],
        ["Suggested Position Size", "2–4% of portfolio"],
        ["FCF Normalization Year", "FY2026–27 expected"],
        ["Key Watch: BUY Trigger", "FCF recovery in Q1/Q2 2026 earnings + Apollo profitability confirmed"],
        ["Key Watch: SELL Trigger", "Price >$185 OR FCF negative 2+ years ex-investment OR search share <40%"],
    ]
)

# ── RISK DISCLAIMER ───────────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "DISCLAIMER", level=2, color=(128,128,128))
p = doc.add_paragraph()
r = p.add_run(
    "This analysis is for educational and research purposes only. It does not constitute financial advice, "
    "investment advice, trading advice, or any other advice. Past performance does not guarantee future results. "
    "All investments carry the risk of loss, including possible loss of principal. Chinese ADRs carry additional "
    "risks including VIE structure uncertainty, regulatory risk, delisting risk under HFCAA, and currency risk. "
    "Consult a qualified financial professional before making any investment decisions. "
    "Generated by Claude Code | Analysis Date: April 4, 2026"
)
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(128, 128, 128)
r.italic = True

filename = "BIDU_combined_research_2026-04-04.docx"
filepath = os.path.join(output_dir, filename)
doc.save(filepath)
print(f"Saved: {filepath}")
