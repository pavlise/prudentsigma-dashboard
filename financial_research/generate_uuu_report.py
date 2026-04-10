from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, datetime

output_dir = r"C:\Users\Pavlos Elpidorou\Documents\AI_Project\financial_research"
os.makedirs(output_dir, exist_ok=True)

TICKER = "UUU"
COMPANY = "Universal Safety Products, Inc."
DATE = datetime.date.today()

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

# ── Helper: set cell background ───────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ── Helper: paragraph style ───────────────────────────────────────────────────
def add_para(text, bold=False, size=11, color=None, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_section_header(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)  # deep navy
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A376C')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(0)
t.paragraph_format.space_after = Pt(2)
r = t.add_run(f"{COMPANY} ({TICKER})")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)
r2 = sub.add_run("Deep Financial Research Report")
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_after = Pt(12)
r3 = date_p.add_run(f"Generated: {DATE.strftime('%B %d, %Y')}  |  Exchange: NYSE American (AMEX)  |  Sector: Consumer Products / Safety")
r3.font.size = Pt(9)
r3.italic = True
r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SNAPSHOT TABLE
# ══════════════════════════════════════════════════════════════════════════════
add_section_header("Market Snapshot (as of April 4, 2026)")

snap_data = [
    ("Current Price", "$5.94"),
    ("Market Capitalization", "~$16.1M"),
    ("Shares Outstanding", "2,717,787"),
    ("52-Week Range", "$1.65 – $8.27"),
    ("12-Month Return", "+232.8%"),
    ("Cash & Equivalents (Sep 30, 2025)", "$5.23M"),
    ("Total Assets (Sep 30, 2025)", "$6.77M"),
    ("Current Ratio", "12.5x"),
    ("Total Debt", "Minimal / None"),
    ("Employees", "~11"),
    ("Analyst Rating (Zacks)", "Neutral"),
]

snap_table = doc.add_table(rows=len(snap_data), cols=2)
snap_table.style = 'Table Grid'
snap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, val) in enumerate(snap_data):
    row = snap_table.rows[i]
    lc = row.cells[0]
    vc = row.cells[1]
    lc.text = label
    vc.text = val
    bg = "EEF2F8" if i % 2 == 0 else "FFFFFF"
    set_cell_bg(lc, bg)
    set_cell_bg(vc, bg)
    lc.paragraphs[0].runs[0].bold = True
    lc.paragraphs[0].runs[0].font.size = Pt(10)
    vc.paragraphs[0].runs[0].font.size = Pt(10)
    lc.width = Inches(2.8)
    vc.width = Inches(2.8)

# ══════════════════════════════════════════════════════════════════════════════
# BUSINESS OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_section_header("Business Overview")
add_para(
    "Universal Safety Products, Inc. (formerly Universal Security Instruments, Inc.) is a Maryland-based "
    "company founded in 1969 that historically designed and distributed smoke alarms, carbon monoxide alarms, "
    "and home safety devices. In May 2025, the company completed a transformative asset sale — divesting its "
    "core smoke and CO alarm business to Feit Electric Company for $6 million in cash. Following this "
    "divestiture, the company retains its wiring devices and bath fans product lines and is actively exploring "
    "new business opportunities. The company now operates with approximately 11 employees and trades on "
    "NYSE American under the ticker UUU.",
    size=10.5
)

# ══════════════════════════════════════════════════════════════════════════════
# FINANCIAL PERFORMANCE
# ══════════════════════════════════════════════════════════════════════════════
add_section_header("Financial Performance")

fin_data = [
    ["Period", "Revenue", "Net Income / (Loss)", "EPS", "Key Driver"],
    ["FY2025 (ended Mar 31, 2025)", "$23.56M", "$500,684", "$0.22", "Full-year alarm business"],
    ["Q4 FY2025 (Jan–Mar 2025)", "$6.23M", "$1.30M", "$0.56", "+40% YoY sales growth"],
    ["H1 FY2026 (Apr–Sep 2025)", "$4.58M", "$810,541", "$0.35", "Includes $2.82M asset sale gain"],
    ["Q2 FY2026 (Jul–Sep 2025)", "$0.76M", "($999,780)", "($0.43)", "Post-divestiture; 89% revenue drop"],
]

fin_table = doc.add_table(rows=len(fin_data), cols=5)
fin_table.style = 'Table Grid'
for i, row_data in enumerate(fin_data):
    row = fin_table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        is_header = (i == 0)
        set_cell_bg(row.cells[j], "1A376C" if is_header else ("F0F4FA" if i % 2 == 0 else "FFFFFF"))
        run = row.cells[j].paragraphs[0].runs
        if run:
            run[0].bold = is_header
            run[0].font.size = Pt(9.5)
            if is_header:
                run[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

add_para(
    "Note: FY2026 operates on an April–March fiscal year. Revenue collapse in Q2 FY2026 reflects the "
    "removal of smoke/CO alarm revenues post-sale. Net income in H1 FY2026 was propped up by the one-time "
    "$2.82M gain from the Feit Electric transaction.",
    size=9.5, color="666666"
)

# ══════════════════════════════════════════════════════════════════════════════
# RECENT DEVELOPMENTS
# ══════════════════════════════════════════════════════════════════════════════
add_section_header("Recent Developments")

devs = [
    ("May 2025 — Core Asset Divestiture:",
     "Feit Electric acquired USI's smoke alarm and CO alarm assets and brand for $6M cash, closing May 22, 2025. "
     "The transaction described as a 'distressed deal' generated a $2.82M accounting gain."),
    ("April 2025 — Corporate Rebrand:",
     "Universal Security Instruments, Inc. officially changed its name to Universal Safety Products, Inc., "
     "reflecting the strategic pivot away from its legacy product identity."),
    ("September 2025 — Special Dividend:",
     "The company declared a one-time $1.00 per share special cash dividend, paid September 25, 2025, "
     "returning ~$2.72M to shareholders. CEO Harvey B. Grossblatt emphasized retaining capital for future growth."),
    ("March 2026 — Aggressive Insider Buying:",
     "Director Milton C. Ault III and affiliated entities (Ault Lending, Hyperscale Data, Alpha Structured Finance) "
     "acquired over 130,000 shares in March 2026 alone, at prices ranging $5.30–$6.35, boosting Ault's "
     "beneficial ownership to 21.1% (585,330 shares)."),
    ("Macro Headwind — Tariff Pressures:",
     "Management cited increased import tariffs across all remaining product lines as a material headwind "
     "to ongoing revenues in the wiring devices and bath fans segments."),
]

for title, text in devs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    r_title = p.add_run(title + " ")
    r_title.bold = True
    r_title.font.size = Pt(10.5)
    r_body = p.add_run(text)
    r_body.font.size = Pt(10.5)

# ══════════════════════════════════════════════════════════════════════════════
# INSIDER TRADING
# ══════════════════════════════════════════════════════════════════════════════
add_section_header("Insider Ownership & Trading Activity")

ins_data = [
    ["Date", "Insider / Entity", "Shares Acquired", "Price Range", "Cumulative Ownership"],
    ["Mar 10–12, 2026", "Ault Lending / Alpha SF / Ault & Co.", "60,500", "$5.30–$5.89", "~21%"],
    ["Mar 2–3, 2026", "Ault-affiliated entities", "27,000", "~$5.40–$5.60", "~20%"],
    ["Mar 19–23, 2026", "Ault Lending, LLC", "24,000", "$6.01–$6.35", "~21%"],
    ["Mar 30–31, 2026", "Ault (direct + affiliated)", "12,608", "Mid-$5 range", "21.1%"],
]

ins_table = doc.add_table(rows=len(ins_data), cols=5)
ins_table.style = 'Table Grid'
for i, row_data in enumerate(ins_data):
    row = ins_table.rows[i]
    for j, ct in enumerate(row_data):
        row.cells[j].text = ct
        is_header = (i == 0)
        set_cell_bg(row.cells[j], "1A376C" if is_header else ("F0F4FA" if i % 2 == 0 else "FFFFFF"))
        runs = row.cells[j].paragraphs[0].runs
        if runs:
            runs[0].bold = is_header
            runs[0].font.size = Pt(9.5)
            if is_header:
                runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

add_para(
    "Milton C. Ault III holds options to purchase 50,000 shares at $3.40 (fully vested Oct 2025, expiring Aug 2035). "
    "Net insider sentiment is overwhelmingly bullish — no recorded insider sells in recent filings.",
    size=9.5, color="444444"
)

# ══════════════════════════════════════════════════════════════════════════════
# VALUATION NOTE (instead of DCF)
# ══════════════════════════════════════════════════════════════════════════════
add_section_header("Valuation Assessment — Cash Shell / Optionality Analysis")
add_para(
    "Given the May 2025 divestiture of UUU's core revenue-generating operations, a conventional DCF analysis "
    "is not meaningful. The company is best assessed as a cash shell with optionality value. Below is a "
    "net asset value (NAV) / liquidation framework:",
    size=10.5
)

nav_data = [
    ["Component", "Estimated Value", "Notes"],
    ["Cash & Equivalents", "~$5.23M", "As of Sep 30, 2025 (post $1/share dividend)"],
    ["Residual Operations (wiring/bath fans)", "~$0.3–0.8M", "Annualizing ~$1.5M rev; operating at loss"],
    ["Inventory & Other Current Assets", "~$0.5–1.0M", "Estimated from total assets of $6.77M"],
    ["Total Estimated NAV", "~$5.5–7.0M", ""],
    ["Current Market Cap", "~$16.1M", "~$5.94 × 2.72M shares"],
    ["Premium to NAV", "~130–190%", "Market pricing in transformation optionality"],
]

nav_table = doc.add_table(rows=len(nav_data), cols=3)
nav_table.style = 'Table Grid'
for i, row_data in enumerate(nav_data):
    row = nav_table.rows[i]
    for j, ct in enumerate(row_data):
        row.cells[j].text = ct
        is_header = (i == 0)
        is_total = (i == 4 or i == 5 or i == 6)
        bg = "1A376C" if is_header else ("FFFDE7" if is_total else ("F0F4FA" if i % 2 == 0 else "FFFFFF"))
        set_cell_bg(row.cells[j], bg)
        runs = row.cells[j].paragraphs[0].runs
        if runs:
            runs[0].bold = is_header or is_total
            runs[0].font.size = Pt(9.5)
            if is_header:
                runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

add_para(
    "Conclusion: At ~$5.94, UUU trades at roughly 1.3–1.9x its estimated net asset value. The market cap "
    "premium (~$9–11M above NAV) reflects speculative optionality — primarily the expectation that Milton "
    "Ault III or management will execute a reverse merger, acquisition, or strategic pivot. This optionality "
    "is highly speculative with no timeline or confirmed strategy.",
    size=10.5
)

# ══════════════════════════════════════════════════════════════════════════════
# BULL / BEAR CASES
# ══════════════════════════════════════════════════════════════════════════════
add_section_header("Bull Case")
bulls = [
    "Cash-rich micro-cap: ~$5.23M cash against $16M market cap — significant balance sheet backing (~33% cash-to-market-cap).",
    "Aggressive insider accumulation: Milton Ault III bought 130K+ shares in March 2026 alone, signalling strong conviction at current prices.",
    "Milton Ault III's track record: Known for engineering corporate transformations; his 21.1% stake and continued buying suggests a strategic play is in motion.",
    "Shell company / reverse merger optionality: With minimal operations and clean books, UUU is an attractive vehicle for a reverse merger or new business acquisition.",
    "No meaningful debt: The company has essentially no financial obligations, giving management maximum flexibility.",
    "Tariff beneficiary potential: A pivot to domestically sourced or manufactured products could benefit from the tariff environment.",
]
for b in bulls:
    add_bullet(b)

add_section_header("Bear Case")
bears = [
    "Revenue collapse: Post-divestiture quarterly revenue of $760K (Q2 FY2026) is barely enough to sustain operations; the core business is gone.",
    "Cash burn: Operating cash flow of -$1.0M means the cash position erodes without a new revenue source.",
    "No articulated strategy: Management has not publicly committed to any specific new business direction — strategic uncertainty is high.",
    "Premium to liquidation value: Market cap is trading at a 130–190% premium to estimated NAV — requires successful pivot to justify.",
    "Extreme illiquidity: With only ~2.72M shares outstanding and 21.1% held by one insider, the free float is tiny and the stock is highly illiquid.",
    "Micro-cap / speculative: Zacks Neutral; no institutional analyst coverage; typical risks of micro-cap speculation apply.",
    "Tariff risk: Remaining wiring devices and bath fans business faces import tariff headwinds.",
]
for b in bears:
    add_bullet(b)

# ══════════════════════════════════════════════════════════════════════════════
# KEY RISKS
# ══════════════════════════════════════════════════════════════════════════════
add_section_header("Key Risks")
risks = [
    "Strategic Execution Risk: The entire investment thesis rests on an undefined future pivot — probability and timeline unknown.",
    "Liquidity Risk: Extremely thin float; large bid/ask spreads and high volatility are expected.",
    "Governance Risk: One individual (Ault) controls 21.1% of voting shares, creating concentration risk.",
    "Regulatory / Tariff Risk: Remaining product lines subject to import tariff escalation.",
    "Cash Depletion Risk: If no new business is identified, cash burn will erode the balance sheet within 3–4 years.",
    "Dilution Risk: A reverse merger or acquisition may involve significant share issuance, diluting existing holders.",
]
for r in risks:
    add_bullet(r)

# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_section_header("Overall Assessment")
add_para(
    "UUU is not a conventional investment — it is a micro-cap cash shell in the early stages of a strategic "
    "transformation. The company sold its 56-year-old core business for $6M in mid-2025, returned capital to "
    "shareholders via a special dividend, and is now effectively a blank-check vehicle with $5M+ in cash and "
    "minimal residual operations. The investment case is almost entirely driven by trust in management "
    "(particularly Milton Ault III) to deploy that capital productively. The aggressive March 2026 insider "
    "buying is the single most compelling positive signal. However, the current market cap of ~$16M implies "
    "the market is already pricing in a successful transformation — leaving limited margin of safety for "
    "investors who enter at these levels without knowing the nature of the pivot.",
    size=10.5
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r_conf = p.add_run("Confidence Level: ")
r_conf.bold = True
r_conf.font.size = Pt(10.5)
r_val = p.add_run("Low-Medium  |  ")
r_val.font.size = Pt(10.5)
r_cat = p.add_run("Category: ")
r_cat.bold = True
r_cat.font.size = Pt(10.5)
r_catv = p.add_run("Speculative / Event-Driven  |  ")
r_catv.font.size = Pt(10.5)
r_rec = p.add_run("Analyst Rating: ")
r_rec.bold = True
r_rec.font.size = Pt(10.5)
r_recv = p.add_run("Neutral (Zacks)")
r_recv.font.size = Pt(10.5)

# ══════════════════════════════════════════════════════════════════════════════
# DISCLAIMER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
disc = doc.add_paragraph()
disc.paragraph_format.space_before = Pt(20)
r_d = disc.add_run(
    "DISCLAIMER: This report is generated for research and informational purposes only and does not constitute "
    "investment advice, a solicitation to buy or sell securities, or a recommendation of any kind. Financial data "
    "sourced from public filings, financial data providers, and news sources. All figures are best-effort estimates "
    "and may not reflect the most current available data. Past performance is not indicative of future results. "
    "Always conduct your own due diligence and consult a qualified financial advisor before making investment decisions."
)
r_d.italic = True
r_d.font.size = Pt(8.5)
r_d.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ──────────────────────────────────────────────────────────────────────
filename = f"{TICKER}_deep_research_{DATE}.docx"
filepath = os.path.join(output_dir, filename)
doc.save(filepath)
print(f"Saved: {filepath}")
