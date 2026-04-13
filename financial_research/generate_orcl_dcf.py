from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, datetime

TICKER = "ORCL"
COMPANY = "Oracle Corporation"
TODAY = datetime.date.today()
output_dir = r"C:\Users\Pavlos Elpidorou\Documents\AI_Project\financial_research"
os.makedirs(output_dir, exist_ok=True)

doc = Document()

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = str(text)
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if bold_first and i == 0:
                cell.paragraphs[0].runs[0].bold = True

def add_header_row(table, headers):
    row = table.rows[0]
    for i, h in enumerate(headers):
        row.cells[i].text = h
        if row.cells[i].paragraphs[0].runs:
            row.cells[i].paragraphs[0].runs[0].bold = True
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

# Title
title = doc.add_heading(f'{COMPANY} ({TICKER}) - DCF Valuation & Intrinsic Value', 0)
title.runs[0].font.size = Pt(18)

doc.add_paragraph(f'Generated: {TODAY.strftime("%B %d, %Y")} | Research Type: DCF Valuation')
doc.add_paragraph('Note: Financial Datasets and Exa MCP servers unavailable. Analysis based on publicly available web data.')
doc.add_paragraph()

# Business Overview
doc.add_heading('1. Business Overview', 1)
doc.add_paragraph(
    'Oracle Corporation is the world\'s largest enterprise database software company, pivoting aggressively '
    'into cloud infrastructure (OCI) and AI services. Revenue for the twelve months ending February 2026 '
    'reached $64.1B (+14.9% YoY), with cloud segment growing 44% YoY in Q3 FY26. The company is targeting '
    '$90B in revenue by FY2027, backed by a $30B+ AI infrastructure fundraising programme.'
)

# Market Data
doc.add_heading('2. Current Price & Market Data', 1)
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_header_row(t, ['Metric', 'Value'])
for row in [
    ('Current Price', '$138.09 (Apr 10, 2026)'),
    ('Market Cap', '$397.2B'),
    ('Shares Outstanding', '2.81B'),
    ('Total Debt', '$95.5B'),
    ('Cash', '$10.8B'),
    ('Net Debt', '$84.7B'),
    ('Enterprise Value (Market Implied)', '$481.9B'),
]:
    add_table_row(t, row, bold_first=True)
doc.add_paragraph()

# FCF History
doc.add_heading('3. FCF History (FY2021-FY2025, Fiscal Year Ends May 31)', 1)
doc.add_paragraph('FY2025 is distorted by a $20B+ AI infrastructure capex surge. Base FCF uses FY2024 ($11.8B).').italic = True
t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
add_header_row(t2, ['Fiscal Year', 'Free Cash Flow', 'Notes'])
for row in [
    ('FY2021', '~$13.5B', 'Pre-Cerner, strong FCF'),
    ('FY2022', '~$9.4B', 'Cerner acquisition capex'),
    ('FY2023', '~$8.5B', 'Integration costs'),
    ('FY2024', '$11.8B', 'Recovery, +39.4% YoY'),
    ('FY2025', '-$0.4B', 'AI infrastructure capex surge (outlier)'),
]:
    add_table_row(t2, row)
doc.add_paragraph()

# WACC
doc.add_heading('4. WACC Estimation', 1)
t3 = doc.add_table(rows=1, cols=3)
t3.style = 'Table Grid'
add_header_row(t3, ['Component', 'Value', 'Notes'])
for row in [
    ('Risk-Free Rate', '4.3%', 'US 10Y Treasury, Apr 2026'),
    ('Beta', '0.96', 'Enterprise software / cloud'),
    ('Equity Risk Premium', '5.0%', 'Damodaran US ERP'),
    ('Cost of Equity', '9.1%', '4.3% + 0.96 x 5.0%'),
    ('Pre-Tax Cost of Debt', '4.55%', 'Weighted avg. bond yield'),
    ('Tax Rate', '18%', 'Oracle effective rate'),
    ('After-Tax Cost of Debt', '3.73%', '4.55% x (1 - 18%)'),
    ('Equity Weight', '80.6%', '$397B / $492.5B'),
    ('Debt Weight', '19.4%', '$95.5B / $492.5B'),
    ('BASE WACC', '8.0%', '80.6% x 9.1% + 19.4% x 3.73%'),
]:
    add_table_row(t3, row, bold_first=True)
doc.add_paragraph()

# DCF Projections
doc.add_heading('5. DCF Projections - Base Case (15% FCF Growth, 5 Years)', 1)
doc.add_paragraph('Base FCF: $11.8B (FY2024). Growth: 15%/yr. WACC: 8.0%. Terminal Growth: 2.5%.')
t4 = doc.add_table(rows=1, cols=4)
t4.style = 'Table Grid'
add_header_row(t4, ['Year', 'FCF ($B)', 'Discount Factor (8%)', 'PV ($B)'])
for row in [
    ('FY2026E', '$13.6B', '0.926', '$12.6B'),
    ('FY2027E', '$15.6B', '0.857', '$13.4B'),
    ('FY2028E', '$17.9B', '0.794', '$14.2B'),
    ('FY2029E', '$20.6B', '0.735', '$15.1B'),
    ('FY2030E', '$23.7B', '0.681', '$16.2B'),
    ('Sum PV FCFs', '', '', '$71.5B'),
    ('Terminal Value (PV)', '2.5% growth', '', '$300.5B'),
    ('Enterprise Value (DCF)', '', '', '$372.0B'),
]:
    add_table_row(t4, row, bold_first=True)
doc.add_paragraph()

# Valuation Summary
doc.add_heading('6. Valuation Summary', 1)
t5 = doc.add_table(rows=1, cols=2)
t5.style = 'Table Grid'
add_header_row(t5, ['Metric', 'Value'])
for row in [
    ('Enterprise Value (DCF)', '$372.0B'),
    ('Less: Net Debt', '($84.7B)'),
    ('Equity Value', '$287.3B'),
    ('Shares Outstanding', '2.81B'),
    ('DCF Fair Value per Share', '$102.2'),
    ('Current Price', '$138.09'),
    ('Upside / (Downside)', '-26.0%'),
    ('VERDICT', 'OVERVALUED'),
    ('Analyst Consensus Target', '$266.23 (35 analysts)'),
]:
    add_table_row(t5, row, bold_first=True)
doc.add_paragraph()

# Sensitivity Analysis
doc.add_heading('7. Sensitivity Analysis (Fair Value per Share, USD)', 1)
doc.add_paragraph('Base FCF: $11.8B | 5-Year Growth: 15% | Net Debt: $84.7B | Shares: 2.81B')
t6 = doc.add_table(rows=4, cols=4)
t6.style = 'Table Grid'
sens_data = [
    ['WACC \\ Terminal Growth', '2.0%', '2.5%', '3.0%'],
    ['7.0%', '$119', '$133', '$151'],
    ['8.0% (Base)', '$93', '$103', '$114'],
    ['9.0%', '$75', '$81', '$89'],
]
for i, row_data in enumerate(sens_data):
    for j, val in enumerate(row_data):
        cell = t6.rows[i].cells[j]
        cell.text = val
        if cell.paragraphs[0].runs:
            if i == 0 or j == 0:
                cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

# Bull/Bear
doc.add_heading('8. Bull Case / Bear Case', 1)
p = doc.add_paragraph()
p.add_run('Bull Case - Target ~$169: ').bold = True
p.add_run('Cloud revenue hits $90B in FY2027; FCF margins recover to 25%+; AI infrastructure investments generate outsized OCI revenues; WACC compresses to 7%.')

p2 = doc.add_paragraph()
p2.add_run('Bear Case - Target ~$75: ').bold = True
p2.add_run('Capex remains elevated beyond FY2027; FCF recovery slower than expected; AWS/Azure/GCP crowd out OCI; rising rates keep WACC at 9%+; $95B debt becomes a burden.')

# Key Risks
doc.add_heading('9. Key Risks', 1)
for risk in [
    'Execution risk on AI infrastructure: $30B+ raised for datacenter buildout must generate returns',
    'Balance sheet leverage: $95.5B total debt, debt-to-equity ratio of 344%',
    'Hyperscaler competition: AWS, Azure, GCP have entrenched enterprise relationships',
    'FCF recovery timeline: If capex stays elevated beyond FY2028, bear case materialises',
    'Valuation gap: 35-analyst consensus of $266 vs. FCF-based DCF of ~$103 reflects speculative premium',
]:
    doc.add_paragraph(risk, style='List Bullet')

# Conclusion
doc.add_heading('10. Conclusion', 1)
doc.add_paragraph(
    'On a traditional FCF-based DCF, Oracle appears overvalued at $138.09, with a base-case fair value of '
    '~$102/share (-26% downside). The stock trades at a ~35% premium to fundamental DCF value, pricing in '
    'a substantial AI-monetisation option that may take 3-5 years to materialise in cash flows. To justify '
    'current prices, Oracle needs near-perfect execution on its hyperscale AI cloud ambitions. Investors '
    'buying today are paying for optionality, not current earnings power.'
)

# Disclaimer
doc.add_paragraph()
p_disc = doc.add_paragraph(
    'Disclaimer: This research is generated by PrudentSigma for informational purposes only. '
    'It does not constitute investment advice. All projections are estimates based on publicly '
    'available data. Past performance is not indicative of future results.'
)
p_disc.runs[0].italic = True
p_disc.runs[0].font.size = Pt(9)

filename = f"{TICKER}_deep_research_{TODAY}.docx"
filepath = os.path.join(output_dir, filename)
doc.save(filepath)
print(f"Saved: {filepath}")
